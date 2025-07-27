#!/usr/bin/env python3
"""
Export service for the Prompt Organizer application.
Supports exporting prompts to multiple formats: PDF, Word, HTML, Markdown, and JSON.
"""

import os
import json
import html
from datetime import datetime
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from enum import Enum


class ExportFormat(Enum):
    """Supported export formats"""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "md"
    JSON = "json"
    TXT = "txt"


@dataclass
class ExportOptions:
    """Export configuration options"""
    format: ExportFormat
    include_metadata: bool = True
    include_tags: bool = True
    include_folders: bool = True
    include_timestamps: bool = True
    include_favorites_only: bool = False
    include_templates_only: bool = False
    custom_title: Optional[str] = None
    custom_description: Optional[str] = None
    group_by_folder: bool = True
    group_by_tags: bool = False
    sort_by: str = "title"  # title, created_at, updated_at
    sort_order: str = "asc"  # asc, desc


@dataclass
class ExportResult:
    """Result of an export operation"""
    success: bool
    file_path: Optional[str] = None
    error_message: Optional[str] = None
    exported_count: int = 0
    file_size: Optional[int] = None


class ExportService:
    """Service for exporting prompts to various formats"""
    
    def __init__(self, database_manager):
        self.db = database_manager
    
    def export_prompts(self, prompt_ids: List[int], output_path: str, 
                      options: ExportOptions) -> ExportResult:
        """Export selected prompts to the specified format"""
        try:
            # Get prompt data
            prompts_data = self._get_prompts_data(prompt_ids, options)
            
            if not prompts_data:
                return ExportResult(
                    success=False,
                    error_message="No prompts found to export"
                )
            
            # Export based on format
            if options.format == ExportFormat.PDF:
                result = self._export_to_pdf(prompts_data, output_path, options)
            elif options.format == ExportFormat.DOCX:
                result = self._export_to_docx(prompts_data, output_path, options)
            elif options.format == ExportFormat.HTML:
                result = self._export_to_html(prompts_data, output_path, options)
            elif options.format == ExportFormat.MARKDOWN:
                result = self._export_to_markdown(prompts_data, output_path, options)
            elif options.format == ExportFormat.JSON:
                result = self._export_to_json(prompts_data, output_path, options)
            elif options.format == ExportFormat.TXT:
                result = self._export_to_txt(prompts_data, output_path, options)
            else:
                return ExportResult(
                    success=False,
                    error_message=f"Unsupported export format: {options.format.value}"
                )
            
            # Add file size if successful
            if result.success and result.file_path and os.path.exists(result.file_path):
                result.file_size = os.path.getsize(result.file_path)
            
            return result
            
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"Export failed: {str(e)}"
            )
    
    def export_all_prompts(self, output_path: str, options: ExportOptions) -> ExportResult:
        """Export all prompts"""
        # Get all prompt IDs
        prompts = self.db.get_prompts()
        prompt_ids = [p['id'] for p in prompts]
        
        # Apply filters
        if options.include_favorites_only:
            prompt_ids = [p['id'] for p in prompts if p['is_favorite']]
        elif options.include_templates_only:
            prompt_ids = [p['id'] for p in prompts if p['is_template']]
        
        return self.export_prompts(prompt_ids, output_path, options)
    
    def export_folder(self, folder_id: int, output_path: str, 
                     options: ExportOptions) -> ExportResult:
        """Export all prompts in a folder"""
        prompts = self.db.get_prompts(folder_id=folder_id)
        prompt_ids = [p['id'] for p in prompts]
        return self.export_prompts(prompt_ids, output_path, options)
    
    def _get_prompts_data(self, prompt_ids: List[int], options: ExportOptions) -> List[Dict]:
        """Get formatted prompt data for export"""
        prompts_data = []
        
        for prompt_id in prompt_ids:
            prompt = self.db.get_prompt(prompt_id)
            if not prompt:
                continue
            
            # Get additional data if needed
            prompt_data = {
                'id': prompt['id'],
                'title': prompt['title'],
                'content': prompt['content'],
                'is_favorite': prompt['is_favorite'],
                'is_template': prompt['is_template'],
                'created_at': prompt['created_at'],
                'updated_at': prompt['updated_at']
            }
            
            if options.include_tags:
                tags = self.db.get_prompt_tags(prompt_id)
                prompt_data['tags'] = [tag['name'] for tag in tags]
            
            if options.include_folders:
                if prompt['folder_id']:
                    folders = self.db.get_all_folders()
                    folder = next((f for f in folders if f['id'] == prompt['folder_id']), None)
                    prompt_data['folder'] = folder['name'] if folder else 'Unknown'
                else:
                    prompt_data['folder'] = 'Root'
            
            prompts_data.append(prompt_data)
        
        # Sort prompts
        reverse = options.sort_order == "desc"
        if options.sort_by == "title":
            prompts_data.sort(key=lambda x: x['title'].lower(), reverse=reverse)
        elif options.sort_by == "created_at":
            prompts_data.sort(key=lambda x: x['created_at'], reverse=reverse)
        elif options.sort_by == "updated_at":
            prompts_data.sort(key=lambda x: x['updated_at'], reverse=reverse)
        
        return prompts_data
    
    def _export_to_html(self, prompts_data: List[Dict], output_path: str, 
                       options: ExportOptions) -> ExportResult:
        """Export to HTML format"""
        try:
            title = options.custom_title or "Prompt Collection"
            description = options.custom_description or "Exported from Prompt Organizer"
            
            html_content = self._generate_html_content(prompts_data, title, description, options)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return ExportResult(
                success=True,
                file_path=output_path,
                exported_count=len(prompts_data)
            )
            
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"HTML export failed: {str(e)}"
            )
    
    def _export_to_markdown(self, prompts_data: List[Dict], output_path: str, 
                           options: ExportOptions) -> ExportResult:
        """Export to Markdown format"""
        try:
            title = options.custom_title or "Prompt Collection"
            description = options.custom_description or "Exported from Prompt Organizer"
            
            md_content = self._generate_markdown_content(prompts_data, title, description, options)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            return ExportResult(
                success=True,
                file_path=output_path,
                exported_count=len(prompts_data)
            )
            
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"Markdown export failed: {str(e)}"
            )
    
    def _export_to_json(self, prompts_data: List[Dict], output_path: str, 
                       options: ExportOptions) -> ExportResult:
        """Export to JSON format"""
        try:
            export_data = {
                'title': options.custom_title or "Prompt Collection",
                'description': options.custom_description or "Exported from Prompt Organizer",
                'exported_at': datetime.now().isoformat(),
                'total_prompts': len(prompts_data),
                'export_options': {
                    'include_metadata': options.include_metadata,
                    'include_tags': options.include_tags,
                    'include_folders': options.include_folders,
                    'group_by_folder': options.group_by_folder,
                    'sort_by': options.sort_by,
                    'sort_order': options.sort_order
                },
                'prompts': prompts_data
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            return ExportResult(
                success=True,
                file_path=output_path,
                exported_count=len(prompts_data)
            )
            
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"JSON export failed: {str(e)}"
            )
    
    def _export_to_txt(self, prompts_data: List[Dict], output_path: str, 
                      options: ExportOptions) -> ExportResult:
        """Export to plain text format"""
        try:
            title = options.custom_title or "Prompt Collection"
            description = options.custom_description or "Exported from Prompt Organizer"
            
            txt_content = self._generate_txt_content(prompts_data, title, description, options)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(txt_content)
            
            return ExportResult(
                success=True,
                file_path=output_path,
                exported_count=len(prompts_data)
            )
            
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"Text export failed: {str(e)}"
            )
    
    def _export_to_pdf(self, prompts_data: List[Dict], output_path: str, 
                      options: ExportOptions) -> ExportResult:
        """Export to PDF format using reportlab"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.darkblue
            )
            
            prompt_title_style = ParagraphStyle(
                'PromptTitle',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.darkgreen
            )
            
            # Add title
            title = options.custom_title or "Prompt Collection"
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Add description
            if options.custom_description:
                story.append(Paragraph(options.custom_description, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Add metadata
            if options.include_metadata:
                metadata = f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>"
                metadata += f"Total prompts: {len(prompts_data)}"
                story.append(Paragraph(metadata, styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Group prompts if needed
            if options.group_by_folder:
                grouped_prompts = self._group_prompts_by_folder(prompts_data)
                for folder_name, folder_prompts in grouped_prompts.items():
                    # Folder header
                    story.append(Paragraph(f"Folder: {folder_name}", styles['Heading2']))
                    story.append(Spacer(1, 12))
                    
                    # Add prompts in folder
                    for prompt in folder_prompts:
                        self._add_prompt_to_pdf_story(prompt, story, styles, prompt_title_style, options)
                    
                    story.append(Spacer(1, 20))
            else:
                # Add all prompts
                for prompt in prompts_data:
                    self._add_prompt_to_pdf_story(prompt, story, styles, prompt_title_style, options)
            
            # Build PDF
            doc.build(story)
            
            return ExportResult(
                success=True,
                file_path=output_path,
                exported_count=len(prompts_data)
            )
            
        except ImportError:
            return ExportResult(
                success=False,
                error_message="PDF export requires reportlab library. Install with: pip install reportlab"
            )
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"PDF export failed: {str(e)}"
            )
    
    def _export_to_docx(self, prompts_data: List[Dict], output_path: str, 
                       options: ExportOptions) -> ExportResult:
        """Export to Word document format using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Add title
            title = options.custom_title or "Prompt Collection"
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add description
            if options.custom_description:
                doc.add_paragraph(options.custom_description)
            
            # Add metadata
            if options.include_metadata:
                metadata_para = doc.add_paragraph()
                metadata_para.add_run(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
                metadata_para.add_run(f"Total prompts: {len(prompts_data)}")
                doc.add_paragraph()  # Empty line
            
            # Group prompts if needed
            if options.group_by_folder:
                grouped_prompts = self._group_prompts_by_folder(prompts_data)
                for folder_name, folder_prompts in grouped_prompts.items():
                    # Folder header
                    doc.add_heading(f"Folder: {folder_name}", level=1)
                    
                    # Add prompts in folder
                    for prompt in folder_prompts:
                        self._add_prompt_to_docx(prompt, doc, options)
            else:
                # Add all prompts
                for prompt in prompts_data:
                    self._add_prompt_to_docx(prompt, doc, options)
            
            # Save document
            doc.save(output_path)
            
            return ExportResult(
                success=True,
                file_path=output_path,
                exported_count=len(prompts_data)
            )
            
        except ImportError:
            return ExportResult(
                success=False,
                error_message="Word export requires python-docx library. Install with: pip install python-docx"
            )
        except Exception as e:
            return ExportResult(
                success=False,
                error_message=f"Word export failed: {str(e)}"
            )
    
    def _generate_html_content(self, prompts_data: List[Dict], title: str, 
                              description: str, options: ExportOptions) -> str:
        """Generate HTML content for export"""
        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "    <meta charset='UTF-8'>",
            "    <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            f"    <title>{html.escape(title)}</title>",
            "    <style>",
            "        body { font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }",
            "        .header { text-align: center; margin-bottom: 40px; }",
            "        .prompt { margin-bottom: 30px; border: 1px solid #ddd; padding: 20px; border-radius: 8px; }",
            "        .prompt-title { color: #2c3e50; margin-bottom: 10px; }",
            "        .prompt-content { background-color: #f8f9fa; padding: 15px; border-radius: 4px; white-space: pre-wrap; }",
            "        .prompt-meta { margin-top: 10px; font-size: 0.9em; color: #666; }",
            "        .tags { margin-top: 10px; }",
            "        .tag { background-color: #007bff; color: white; padding: 2px 8px; border-radius: 12px; font-size: 0.8em; margin-right: 5px; }",
            "        .folder-section { margin-bottom: 40px; }",
            "        .folder-title { color: #e74c3c; border-bottom: 2px solid #e74c3c; padding-bottom: 5px; }",
            "        .favorite { border-left: 4px solid #f39c12; }",
            "        .template { border-left: 4px solid #9b59b6; }",
            "    </style>",
            "</head>",
            "<body>",
            "    <div class='header'>",
            f"        <h1>{html.escape(title)}</h1>",
            f"        <p>{html.escape(description)}</p>",
        ]
        
        if options.include_metadata:
            html_parts.extend([
                f"        <p><strong>Exported on:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>",
                f"        <p><strong>Total prompts:</strong> {len(prompts_data)}</p>",
            ])
        
        html_parts.append("    </div>")
        
        # Group prompts if needed
        if options.group_by_folder:
            grouped_prompts = self._group_prompts_by_folder(prompts_data)
            for folder_name, folder_prompts in grouped_prompts.items():
                html_parts.extend([
                    "    <div class='folder-section'>",
                    f"        <h2 class='folder-title'>Folder: {html.escape(folder_name)}</h2>",
                ])
                
                for prompt in folder_prompts:
                    html_parts.append(self._generate_html_prompt(prompt, options))
                
                html_parts.append("    </div>")
        else:
            for prompt in prompts_data:
                html_parts.append(self._generate_html_prompt(prompt, options))
        
        html_parts.extend([
            "</body>",
            "</html>"
        ])
        
        return "\n".join(html_parts)
    
    def _generate_html_prompt(self, prompt: Dict, options: ExportOptions) -> str:
        """Generate HTML for a single prompt"""
        css_classes = ["prompt"]
        if prompt.get('is_favorite'):
            css_classes.append("favorite")
        if prompt.get('is_template'):
            css_classes.append("template")
        
        html_parts = [
            f"    <div class='{' '.join(css_classes)}'>",
            f"        <h3 class='prompt-title'>{html.escape(prompt['title'])}</h3>",
            f"        <div class='prompt-content'>{html.escape(prompt['content'])}</div>",
        ]
        
        if options.include_metadata or options.include_tags:
            meta_parts = []
            
            if options.include_metadata:
                if options.include_timestamps:
                    meta_parts.append(f"Created: {prompt['created_at']}")
                    meta_parts.append(f"Updated: {prompt['updated_at']}")
                
                if prompt.get('is_favorite'):
                    meta_parts.append("⭐ Favorite")
                if prompt.get('is_template'):
                    meta_parts.append("📋 Template")
            
            if meta_parts:
                html_parts.append(f"        <div class='prompt-meta'>{' | '.join(meta_parts)}</div>")
            
            if options.include_tags and prompt.get('tags'):
                tags_html = ''.join([f"<span class='tag'>{html.escape(tag)}</span>" for tag in prompt['tags']])
                html_parts.append(f"        <div class='tags'>{tags_html}</div>")
        
        html_parts.append("    </div>")
        return "\n".join(html_parts)
    
    def _generate_markdown_content(self, prompts_data: List[Dict], title: str, 
                                  description: str, options: ExportOptions) -> str:
        """Generate Markdown content for export"""
        md_parts = [
            f"# {title}",
            "",
            description,
            ""
        ]
        
        if options.include_metadata:
            md_parts.extend([
                f"**Exported on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ",
                f"**Total prompts:** {len(prompts_data)}",
                ""
            ])
        
        # Group prompts if needed
        if options.group_by_folder:
            grouped_prompts = self._group_prompts_by_folder(prompts_data)
            for folder_name, folder_prompts in grouped_prompts.items():
                md_parts.extend([
                    f"## Folder: {folder_name}",
                    ""
                ])
                
                for prompt in folder_prompts:
                    md_parts.extend(self._generate_markdown_prompt(prompt, options))
                    md_parts.append("")
        else:
            for prompt in prompts_data:
                md_parts.extend(self._generate_markdown_prompt(prompt, options))
                md_parts.append("")
        
        return "\n".join(md_parts)
    
    def _generate_markdown_prompt(self, prompt: Dict, options: ExportOptions) -> List[str]:
        """Generate Markdown for a single prompt"""
        md_parts = [f"### {prompt['title']}"]
        
        # Add badges
        badges = []
        if prompt.get('is_favorite'):
            badges.append("⭐ Favorite")
        if prompt.get('is_template'):
            badges.append("📋 Template")
        
        if badges:
            md_parts.append(f"*{' | '.join(badges)}*")
        
        md_parts.extend([
            "",
            "```",
            prompt['content'],
            "```"
        ])
        
        if options.include_metadata or options.include_tags:
            meta_parts = []
            
            if options.include_metadata and options.include_timestamps:
                meta_parts.extend([
                    f"**Created:** {prompt['created_at']}",
                    f"**Updated:** {prompt['updated_at']}"
                ])
            
            if options.include_tags and prompt.get('tags'):
                tags_str = ", ".join([f"`{tag}`" for tag in prompt['tags']])
                meta_parts.append(f"**Tags:** {tags_str}")
            
            if meta_parts:
                md_parts.extend([""] + meta_parts)
        
        return md_parts
    
    def _generate_txt_content(self, prompts_data: List[Dict], title: str, 
                             description: str, options: ExportOptions) -> str:
        """Generate plain text content for export"""
        txt_parts = [
            "=" * len(title),
            title,
            "=" * len(title),
            "",
            description,
            ""
        ]
        
        if options.include_metadata:
            txt_parts.extend([
                f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                f"Total prompts: {len(prompts_data)}",
                ""
            ])
        
        # Group prompts if needed
        if options.group_by_folder:
            grouped_prompts = self._group_prompts_by_folder(prompts_data)
            for folder_name, folder_prompts in grouped_prompts.items():
                txt_parts.extend([
                    f"FOLDER: {folder_name}",
                    "-" * (8 + len(folder_name)),
                    ""
                ])
                
                for prompt in folder_prompts:
                    txt_parts.extend(self._generate_txt_prompt(prompt, options))
                    txt_parts.append("")
        else:
            for prompt in prompts_data:
                txt_parts.extend(self._generate_txt_prompt(prompt, options))
                txt_parts.append("")
        
        return "\n".join(txt_parts)
    
    def _generate_txt_prompt(self, prompt: Dict, options: ExportOptions) -> List[str]:
        """Generate plain text for a single prompt"""
        txt_parts = [
            prompt['title'],
            "-" * len(prompt['title'])
        ]
        
        # Add badges
        badges = []
        if prompt.get('is_favorite'):
            badges.append("[FAVORITE]")
        if prompt.get('is_template'):
            badges.append("[TEMPLATE]")
        
        if badges:
            txt_parts.append(" ".join(badges))
        
        txt_parts.extend([
            "",
            prompt['content']
        ])
        
        if options.include_metadata or options.include_tags:
            txt_parts.append("")
            
            if options.include_metadata and options.include_timestamps:
                txt_parts.extend([
                    f"Created: {prompt['created_at']}",
                    f"Updated: {prompt['updated_at']}"
                ])
            
            if options.include_tags and prompt.get('tags'):
                tags_str = ", ".join(prompt['tags'])
                txt_parts.append(f"Tags: {tags_str}")
        
        return txt_parts
    
    def _group_prompts_by_folder(self, prompts_data: List[Dict]) -> Dict[str, List[Dict]]:
        """Group prompts by folder"""
        grouped = {}
        for prompt in prompts_data:
            folder_name = prompt.get('folder', 'Root')
            if folder_name not in grouped:
                grouped[folder_name] = []
            grouped[folder_name].append(prompt)
        return grouped
    
    def _add_prompt_to_pdf_story(self, prompt: Dict, story: List, styles: Any, 
                                prompt_title_style: Any, options: ExportOptions):
        """Add a prompt to PDF story"""
        from reportlab.platypus import Paragraph, Spacer
        
        # Title with badges
        title_text = prompt['title']
        if prompt.get('is_favorite'):
            title_text += " ⭐"
        if prompt.get('is_template'):
            title_text += " 📋"
        
        story.append(Paragraph(title_text, prompt_title_style))
        story.append(Spacer(1, 6))
        
        # Content
        content_text = prompt['content'].replace('\n', '<br/>')
        story.append(Paragraph(content_text, styles['Normal']))
        
        # Metadata
        if options.include_metadata or options.include_tags:
            meta_parts = []
            
            if options.include_metadata and options.include_timestamps:
                meta_parts.append(f"Created: {prompt['created_at']}")
                meta_parts.append(f"Updated: {prompt['updated_at']}")
            
            if options.include_tags and prompt.get('tags'):
                tags_str = ", ".join(prompt['tags'])
                meta_parts.append(f"Tags: {tags_str}")
            
            if meta_parts:
                meta_text = " | ".join(meta_parts)
                story.append(Spacer(1, 6))
                story.append(Paragraph(f"<i>{meta_text}</i>", styles['Normal']))
        
        story.append(Spacer(1, 20))
    
    def _add_prompt_to_docx(self, prompt: Dict, doc: Any, options: ExportOptions):
        """Add a prompt to Word document"""
        # Title
        title_text = prompt['title']
        if prompt.get('is_favorite'):
            title_text += " ⭐"
        if prompt.get('is_template'):
            title_text += " 📋"
        
        doc.add_heading(title_text, level=2)
        
        # Content
        doc.add_paragraph(prompt['content'])
        
        # Metadata
        if options.include_metadata or options.include_tags:
            meta_parts = []
            
            if options.include_metadata and options.include_timestamps:
                meta_parts.append(f"Created: {prompt['created_at']}")
                meta_parts.append(f"Updated: {prompt['updated_at']}")
            
            if options.include_tags and prompt.get('tags'):
                tags_str = ", ".join(prompt['tags'])
                meta_parts.append(f"Tags: {tags_str}")
            
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(" | ".join(meta_parts))
                meta_run.italic = True
        
        doc.add_paragraph()  # Empty line
    
    def get_supported_formats(self) -> List[ExportFormat]:
        """Get list of supported export formats"""
        supported = [ExportFormat.HTML, ExportFormat.MARKDOWN, ExportFormat.JSON, ExportFormat.TXT]
        
        # Check for optional dependencies
        try:
            import reportlab
            supported.append(ExportFormat.PDF)
        except ImportError:
            pass
        
        try:
            import docx
            supported.append(ExportFormat.DOCX)
        except ImportError:
            pass
        
        return supported
    
    def get_format_extension(self, format: ExportFormat) -> str:
        """Get file extension for format"""
        extensions = {
            ExportFormat.PDF: ".pdf",
            ExportFormat.DOCX: ".docx",
            ExportFormat.HTML: ".html",
            ExportFormat.MARKDOWN: ".md",
            ExportFormat.JSON: ".json",
            ExportFormat.TXT: ".txt"
        }
        return extensions.get(format, ".txt")
    
    def validate_export_options(self, options: ExportOptions) -> List[str]:
        """Validate export options and return list of issues"""
        issues = []
        
        if options.format not in self.get_supported_formats():
            issues.append(f"Export format {options.format.value} is not supported or requires additional dependencies")
        
        if options.sort_by not in ["title", "created_at", "updated_at"]:
            issues.append(f"Invalid sort_by option: {options.sort_by}")
        
        if options.sort_order not in ["asc", "desc"]:
            issues.append(f"Invalid sort_order option: {options.sort_order}")
        
        return issues